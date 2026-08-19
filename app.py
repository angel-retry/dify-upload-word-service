import os
import requests
import json
import uuid
import re
import gc
import platform
import subprocess
import asyncio
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
import docx
from docx.oxml.ns import qn

load_dotenv()

app = FastAPI()

BASE_DIR = Path(__file__).resolve().parent

# 精簡為兩個資料夾：
# - docx：處理後的版本（縮排修正＋自動編號硬寫成文字），同時拿去轉 PDF 跟餵給 Dify，
#         只保留這一份，不另外保存使用者上傳的未處理原始檔
# - pdf：僅供瀏覽用的轉檔結果（PDF 網址已寫入 Dify metadata，不需要另存注入 PDF 說明段落的 Word 檔）
STORAGE_DIR = BASE_DIR / "storage"
DOCX_DIR = STORAGE_DIR / "docx"
PDF_DIR = STORAGE_DIR / "pdf"

DOCX_DIR.mkdir(parents=True, exist_ok=True)
PDF_DIR.mkdir(parents=True, exist_ok=True)

# 僅掛載 PDF 目錄為靜態路由，確保外部無法存取原始 Word 檔
app.mount("/static/pdf", StaticFiles(directory=str(PDF_DIR)), name="static_pdf")

SERVER_BASE_URL = os.getenv("SERVER_BASE_URL", "http://localhost:8000")
DIFY_API_BASE = os.getenv("DIFY_API_BASE", "http://localhost/v1")
DIFY_API_KEY = os.getenv("DIFY_API_KEY")

if not DIFY_API_KEY:
    raise ValueError("未設定 DIFY_API_KEY，請確認 .env 檔案設定。")

# metadata 欄位 ID 的記憶體快取，不用再手動把 ID 貼進 .env。
# 第一次用到某個欄位名稱時才會打一次查詢 API，之後同一次程式執行期間直接沿用快取，
# 不會每次上傳都重打請求。若欄位被刪除重建、快取失效，會自動重新查詢。
_METADATA_FIELD_CACHE = {}


def get_metadata_field_id(dataset_id: str, field_name: str, headers: dict, force_refresh: bool = False) -> str:
    """
    動態取得某個 metadata 欄位（例如 file_uuid、pdf_url）在 Dify 裡的欄位 ID。
    優先從記憶體快取拿；快取沒有時才打 API 查詢一次並存進快取。
    快取 key 改用 (dataset_id, field_name)，避免不同知識庫的欄位 ID 互相污染。
    為避免觸發 Dify 1.16.0-rc1 資料庫寫入異常，已移除自動建立欄位的邏輯。
    """
    cache_key = (dataset_id, field_name)
    if not force_refresh and cache_key in _METADATA_FIELD_CACHE:
        return _METADATA_FIELD_CACHE[cache_key]

    print(f"[Metadata查詢] 快取中沒有 '{field_name}'，向 Dify 查詢目前所有欄位...")
    url = f"{DIFY_API_BASE}/datasets/{dataset_id}/metadata"
    try:
        resp = requests.get(url, headers=headers)
        if resp.status_code == 200:
            fields = resp.json().get("doc_metadata", [])
            for f in fields:
                if f.get("name") == field_name:
                    _METADATA_FIELD_CACHE[cache_key] = f.get("id")
                    print(f"[Metadata查詢] 找到欄位 '{field_name}'，ID={f.get('id')}，已存入快取。")
                    return f.get("id")
        else:
            print(f"[Metadata查詢] 查詢欄位列表失敗，狀態碼: {resp.status_code}")
    except Exception as e:
        print(f"[Metadata查詢] 發生例外錯誤: {str(e)}")

    print(f"[Metadata查詢] 欄位 '{field_name}' 不存在，為避免資料庫異常，不執行自動建立。")
    return None

LIBREOFFICE_WIN_PATH = os.getenv(
    "LIBREOFFICE_WIN_PATH",
    r"C:\Program Files\LibreOffice\program\soffice.exe"
)


def safe_unlink(file_path: Path):
    if file_path.exists():
        try:
            file_path.unlink()
            print(f"[檔案刪除] 已成功刪除本地檔案: {file_path}")
        except Exception as e:
            print(f"[檔案刪除] 刪除本地檔案失敗 ({file_path}): {str(e)}")


def set_document_metadata(dataset_id: str, document_id: str, file_uuid: str, pdf_url: str, headers: dict) -> bool:
    """
    把 file_uuid 與 pdf_url 一起寫進該 Dify 文件的 metadata（合併成一次請求）。
    欄位 ID 改為動態查詢＋快取（見 get_metadata_field_id），不需要手動在 .env 設定 ID，
    也不怕欄位被刪除重建後 ID 對不上的問題。

    注意：這支端點官方文件目前未列出（屬於社群驗證可用、但未正式收錄的端點），
    不同 Dify 版本行為可能有落差，建議先用 curl/Postman 對你的環境測試過一次。
    """
    field_ids = {
        "file_uuid": get_metadata_field_id(dataset_id, "file_uuid", headers),
        "pdf_url": get_metadata_field_id(dataset_id, "pdf_url", headers),
    }
    values = {"file_uuid": file_uuid, "pdf_url": pdf_url}

    metadata_list = [
        {"id": field_ids[name], "name": name, "value": values[name]}
        for name in ("file_uuid", "pdf_url")
        if field_ids[name]
    ]

    if not metadata_list:
        print("[Metadata寫入] 找不到任何可用的 metadata 欄位 ID，略過寫入。")
        return False

    def _do_write(list_to_write):
        url = f"{DIFY_API_BASE}/datasets/{dataset_id}/documents/metadata"
        payload = {
            "operation_data": [
                {
                    "document_id": document_id,
                    "metadata_list": list_to_write,
                    "partial_update": True
                }
            ]
        }
        return requests.post(url, headers=headers, json=payload)

    try:
        resp = _do_write(metadata_list)
        if resp.status_code in (200, 202):
            print(f"[Metadata寫入] 成功將 file_uuid={file_uuid}, pdf_url={pdf_url} 寫入文件 {document_id}")
            return True

        # 若失敗，可能是快取裡的欄位 ID 已經過期（欄位被刪除重建過），
        # 強制重新查詢一次欄位 ID 後再重試一次。
        print(f"[Metadata寫入] 第一次嘗試失敗，狀態碼: {resp.status_code}，內容: {resp.text}")
        print("[Metadata寫入] 懷疑快取的欄位 ID 已過期，強制重新查詢後重試一次...")
        field_ids = {
            "file_uuid": get_metadata_field_id(dataset_id, "file_uuid", headers, force_refresh=True),
            "pdf_url": get_metadata_field_id(dataset_id, "pdf_url", headers, force_refresh=True),
        }
        metadata_list = [
            {"id": field_ids[name], "name": name, "value": values[name]}
            for name in ("file_uuid", "pdf_url")
            if field_ids[name]
        ]
        if not metadata_list:
            return False

        resp2 = _do_write(metadata_list)
        if resp2.status_code in (200, 202):
            print(f"[Metadata寫入] 重試成功，已寫入文件 {document_id}")
            return True
        print(f"[Metadata寫入] 重試仍失敗，狀態碼: {resp2.status_code}，內容: {resp2.text}")
        return False
    except Exception as e:
        print(f"[Metadata寫入] 發生例外錯誤: {str(e)}")
        return False


def get_file_uuid_from_metadata(doc: dict) -> str:
    """
    嘗試從「取得文件列表」API 回傳的單筆 doc 物件中，
    解析出 file_uuid metadata 的值。
    Dify 的 doc_metadata 欄位常見格式為 list，例如：
    [{"id": "...", "name": "file_uuid", "value": "xxxx-xxxx"}]
    若你的 Dify 版本格式不同，請依實際回傳結構調整這裡。
    """
    doc_metadata = doc.get("doc_metadata") or doc.get("metadata") or []
    if isinstance(doc_metadata, list):
        for item in doc_metadata:
            if isinstance(item, dict) and item.get("name") == "file_uuid" and item.get("value"):
                return item["value"]
    elif isinstance(doc_metadata, dict):
        if doc_metadata.get("file_uuid"):
            return doc_metadata["file_uuid"]
    return None


# ============================================================
# 取得段落「實際生效」的自動編號 (numId / ilvl)
#
# 問題背景：
# Word 的自動編號清單有兩種掛法：
#   1) 編號直接掛在段落自己的 w:pPr/w:numPr 上
#   2) 編號掛在該段落套用的「樣式」上（styles.xml 裡該 style 的
#      w:pPr/w:numPr），段落本身完全沒有 w:numPr，是靠套用樣式
#      繼承來的編號
# 舊版程式只檢查第 1 種情況，導致像「內文-數字標示」這類靠樣式
# 繼承編號的段落，自動編號完全沒被抓到、也就沒被硬寫成文字。
# ============================================================
def _get_style_numpr(doc: docx.Document, style_id: str, visited=None):
    """
    往上查某個 style 是否有 w:numPr；若自己沒有，沿著 w:basedOn 鏈
    繼續往上層樣式找（Word 的樣式繼承機制）。
    回傳 (numId, ilvl) 或 None。
    """
    if visited is None:
        visited = set()
    if not style_id or style_id in visited:
        return None
    visited.add(style_id)

    for s in doc.styles.element.findall(qn('w:style')):
        if s.get(qn('w:styleId')) == style_id:
            pPr = s.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    ilvl_el = numPr.find(qn('w:ilvl'))
                    numId_el = numPr.find(qn('w:numId'))
                    ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'
                    num_id = numId_el.get(qn('w:val')) if numId_el is not None else None
                    if num_id is not None:
                        return (num_id, ilvl)
                basedOn = pPr.find(qn('w:basedOn'))
                if basedOn is not None:
                    return _get_style_numpr(doc, basedOn.get(qn('w:val')), visited)
            return None
    return None


def get_effective_numpr(doc: docx.Document, p):
    """
    取得段落「實際生效」的編號來源，優先序：
      段落自己的 w:numPr  >  段落套用樣式（含 basedOn 繼承鏈）的 w:numPr
    回傳 (num_id, ilvl, is_direct)，is_direct=True 代表編號是直接掛在
    段落自己身上（段落的 w:numPr 節點可以直接刪除來取消編號）；
    is_direct=False 代表編號是樣式繼承來的（段落本身沒有 w:numPr 節點
    可刪，必須額外「明確覆蓋成 numId=0」才能取消繼承的編號）。
    查無任何編號（含 numId=0 代表本來就已停用）則回傳 None。
    """
    pPr = p._p.pPr
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl_el = numPr.find(qn('w:ilvl'))
            numId_el = numPr.find(qn('w:numId'))
            ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'
            num_id = numId_el.get(qn('w:val')) if numId_el is not None else None
            if num_id is not None and num_id != '0':
                return (num_id, ilvl, True)
            if num_id == '0':
                return None  # 段落自己明確關閉編號，不該再往樣式查

    style_id = None
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        style_id = pStyle.get(qn('w:val')) if pStyle is not None else None
    if style_id is None and p.style is not None:
        style_id = p.style.style_id

    style_result = _get_style_numpr(doc, style_id)
    if style_result is not None:
        num_id, ilvl = style_result
        if num_id != '0':
            return (num_id, ilvl, False)
    return None


# ============================================================
# 修正 w:leftChars 造成的 LibreOffice 縮排跑版問題（完全保留未動）
#
# 執行順序很重要：這個函式一定要在 normalize_and_solidify_list_numbering
# 之前執行。因為這裡會把每個編號段落該有的 left / hanging 縮排值，
# 直接寫死到段落自己的 w:ind 上（不依賴 w:numPr 是否還存在）；
# 之後 normalize_and_solidify_list_numbering 把 w:numPr 移除時，
# 縮排效果不會跟著消失，排版還是正確的。
# ============================================================
def fix_list_indentation_for_libreoffice(doc: docx.Document):
    """
    問題根源：
    段落 pPr 裡常見 <w:ind w:leftChars="0"/>（以字元為單位的縮排覆蓋值）。
    Word 能正確運算這個值疊加 numbering.xml 定義的階層縮排，
    但 LibreOffice 對 leftChars 的相容性較差，常常直接把它當成
    「縮排=0」，導致所有層級（1. / 1.1 / 1.1.1）全部貼齊左邊界，
    階層感消失。

    解法：
    完全不動 w:numPr（自動編號機制維持原樣，不會有編號重複問題），
    只針對每個有編號的段落，去查 numbering.xml 中對應
    numId + ilvl（或 lvlOverride）該有的 left / hanging 縮排值，
    直接寫死覆蓋到段落自己的 w:ind 上，並移除會讓 LibreOffice
    誤判的 leftChars / rightChars 屬性。
    """
    numbering_part = doc.part.numbering_part
    numbering_root = numbering_part.element

    # 1) numId -> abstractNumId 對照表（含 lvlOverride 的獨立 ind 設定）
    num_to_abstract = {}
    num_lvl_overrides = {}  # numId -> {ilvl: (left, hanging)}

    for num_el in numbering_root.findall(qn('w:num')):
        num_id = num_el.get(qn('w:numId'))
        abstract_el = num_el.find(qn('w:abstractNumId'))
        if abstract_el is not None:
            num_to_abstract[num_id] = abstract_el.get(qn('w:val'))

        overrides = {}
        for lvl_override in num_el.findall(qn('w:lvlOverride')):
            ilvl = lvl_override.get(qn('w:ilvl'))
            lvl_el = lvl_override.find(qn('w:lvl'))
            if lvl_el is not None:
                pPr_el = lvl_el.find(qn('w:pPr'))
                ind_el = pPr_el.find(qn('w:ind')) if pPr_el is not None else None
                if ind_el is not None:
                    left = ind_el.get(qn('w:left'))
                    hanging = ind_el.get(qn('w:hanging'))
                    if left is not None:
                        overrides[ilvl] = (left, hanging)
        if overrides:
            num_lvl_overrides[num_id] = overrides

    # 2) abstractNumId -> {ilvl: (left, hanging)} 對照表
    abstract_lvl_ind = {}
    for abstract_el in numbering_root.findall(qn('w:abstractNum')):
        abstract_id = abstract_el.get(qn('w:abstractNumId'))
        lvl_map = {}
        for lvl_el in abstract_el.findall(qn('w:lvl')):
            ilvl = lvl_el.get(qn('w:ilvl'))
            pPr_el = lvl_el.find(qn('w:pPr'))
            ind_el = pPr_el.find(qn('w:ind')) if pPr_el is not None else None
            if ind_el is not None:
                left = ind_el.get(qn('w:left'))
                hanging = ind_el.get(qn('w:hanging'))
                if left is not None:
                    lvl_map[ilvl] = (left, hanging)
        abstract_lvl_ind[abstract_id] = lvl_map

    fixed_count = 0

    for p in doc.paragraphs:
        # 改用 get_effective_numpr：不只看段落自己的 w:numPr，
        # 也涵蓋編號掛在「樣式」上、段落靠套用樣式繼承編號的情況。
        result = get_effective_numpr(doc, p)
        if result is None:
            continue  # 沒有自動編號（無論直接掛還是樣式繼承），跳過，不動它的縮排
        num_id, ilvl, is_direct = result

        pPr = p._p.get_or_add_pPr()

        # 優先找 lvlOverride，其次找 abstractNum 預設層級縮排
        left = hanging = None
        if num_id in num_lvl_overrides and ilvl in num_lvl_overrides[num_id]:
            left, hanging = num_lvl_overrides[num_id][ilvl]
        else:
            abstract_id = num_to_abstract.get(num_id)
            if abstract_id is not None:
                lvl_map = abstract_lvl_ind.get(abstract_id, {})
                if ilvl in lvl_map:
                    left, hanging = lvl_map[ilvl]

        if left is None:
            continue  # 查不到對照值，保守跳過，避免誤改

        # 3) 覆蓋段落自己的 w:ind：寫死 left / hanging（twips），
        #    並移除 leftChars / rightChars，避免 LibreOffice 誤判
        ind_el = pPr.find(qn('w:ind'))
        if ind_el is None:
            ind_el = pPr.makeelement(qn('w:ind'), {})
            pPr.append(ind_el)

        ind_el.set(qn('w:left'), left)
        if hanging is not None:
            ind_el.set(qn('w:hanging'), hanging)

        for attr in ('w:leftChars', 'w:rightChars', 'w:hangingChars', 'w:firstLineChars'):
            if ind_el.get(qn(attr)) is not None:
                del ind_el.attrib[qn(attr)]

        fixed_count += 1

    print(f"[縮排修正] 共修正 {fixed_count} 個段落的縮排數值")


_ROMAN_TABLE = [
    (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
    (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
    (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I'),
]

_CHINESE_DIGITS = "〇一二三四五六七八九"
_HEAVENLY_STEMS = "甲乙丙丁戊己庚辛壬癸"  # 天干，Word 的甲乙丙清單樣式常用這十個字循環


def _to_roman(num: int) -> str:
    result = []
    for value, symbol in _ROMAN_TABLE:
        while num >= value:
            result.append(symbol)
            num -= value
    return "".join(result)


def _to_letters(num: int) -> str:
    """1 -> a, 2 -> b, ... 26 -> z, 27 -> aa, 28 -> ab ..."""
    letters = ""
    while num > 0:
        num, remainder = divmod(num - 1, 26)
        letters = chr(97 + remainder) + letters
    return letters


def _to_chinese_numeral(num: int) -> str:
    """簡易中文數字轉換，涵蓋清單常見的 1~99 範圍"""
    if num <= 0:
        return str(num)
    if num < 10:
        return _CHINESE_DIGITS[num]
    if num < 20:
        ones = num - 10
        return "十" + (_CHINESE_DIGITS[ones] if ones else "")
    tens, ones = divmod(num, 10)
    result = _CHINESE_DIGITS[tens] + "十"
    if ones:
        result += _CHINESE_DIGITS[ones]
    return result


def _to_heavenly_stem(num: int) -> str:
    """甲乙丙丁...，超過十個循環使用（天干本身只有十個字）"""
    if num <= 0:
        return str(num)
    return _HEAVENLY_STEMS[(num - 1) % len(_HEAVENLY_STEMS)]


def _format_counter(num: int, num_fmt: str) -> str:
    """
    依 Word 的 w:numFmt 值，把計數轉成對應樣式的文字。
    涵蓋常見樣式：阿拉伯數字、a/b/c、A/B/C、i/ii/iii、I/II/III、
    一/二/三（中文計數）、甲/乙/丙（天干）。
    遇到沒對應到的冷門格式，保守 fallback 成阿拉伯數字，並印出提醒。
    """
    if num_fmt in ("decimal", "decimalHalfWidth"):
        return str(num)
    if num_fmt == "decimalZero":
        return f"{num:02d}"
    if num_fmt == "lowerLetter":
        return _to_letters(num)
    if num_fmt == "upperLetter":
        return _to_letters(num).upper()
    if num_fmt == "lowerRoman":
        return _to_roman(num).lower()
    if num_fmt == "upperRoman":
        return _to_roman(num)
    if num_fmt in ("chineseCounting", "chineseCountingThousand", "ideographDigital"):
        return _to_chinese_numeral(num)
    if num_fmt in (
        "taiwaneseCounting", "taiwaneseCountingThousand",
        "ideographTraditional", "ideographZodiac", "ideographZodiacTraditional",
        "ideographLegalTraditional",
    ):
        return _to_heavenly_stem(num)

    print(f"[編號硬寫] 遇到未特別支援的清單格式 '{num_fmt}'，暫以阿拉伯數字代替。")
    return str(num)


def _build_numbering_level_defs(doc: docx.Document):
    """
    掃過 numbering.xml，建立 (numId, ilvl) -> (numFmt, lvlText) 的對照表。
    lvlOverride（若有自訂 numFmt / lvlText）優先於 abstractNum 的預設定義。
    """
    numbering_root = doc.part.numbering_part.element

    num_to_abstract = {}
    for num_el in numbering_root.findall(qn('w:num')):
        num_id = num_el.get(qn('w:numId'))
        abstract_el = num_el.find(qn('w:abstractNumId'))
        if abstract_el is not None:
            num_to_abstract[num_id] = abstract_el.get(qn('w:val'))

    abstract_lvl_defs = {}
    for abstract_el in numbering_root.findall(qn('w:abstractNum')):
        abstract_id = abstract_el.get(qn('w:abstractNumId'))
        lvl_map = {}
        for lvl_el in abstract_el.findall(qn('w:lvl')):
            ilvl = lvl_el.get(qn('w:ilvl'))
            num_fmt_el = lvl_el.find(qn('w:numFmt'))
            lvl_text_el = lvl_el.find(qn('w:lvlText'))
            num_fmt = num_fmt_el.get(qn('w:val')) if num_fmt_el is not None else "decimal"
            lvl_text = lvl_text_el.get(qn('w:val')) if lvl_text_el is not None else "%1."
            lvl_map[ilvl] = (num_fmt, lvl_text)
        abstract_lvl_defs[abstract_id] = lvl_map

    num_lvl_overrides = {}
    for num_el in numbering_root.findall(qn('w:num')):
        num_id = num_el.get(qn('w:numId'))
        overrides = {}
        for lvl_override in num_el.findall(qn('w:lvlOverride')):
            ilvl = lvl_override.get(qn('w:ilvl'))
            lvl_el = lvl_override.find(qn('w:lvl'))
            if lvl_el is not None:
                num_fmt_el = lvl_el.find(qn('w:numFmt'))
                lvl_text_el = lvl_el.find(qn('w:lvlText'))
                if num_fmt_el is not None or lvl_text_el is not None:
                    abstract_id = num_to_abstract.get(num_id)
                    default_fmt, default_text = abstract_lvl_defs.get(abstract_id, {}).get(ilvl, ("decimal", "%1."))
                    num_fmt = num_fmt_el.get(qn('w:val')) if num_fmt_el is not None else default_fmt
                    lvl_text = lvl_text_el.get(qn('w:val')) if lvl_text_el is not None else default_text
                    overrides[ilvl] = (num_fmt, lvl_text)
        if overrides:
            num_lvl_overrides[num_id] = overrides

    def get_level_def(num_id: str, ilvl: str):
        if num_id in num_lvl_overrides and ilvl in num_lvl_overrides[num_id]:
            return num_lvl_overrides[num_id][ilvl]
        abstract_id = num_to_abstract.get(num_id)
        if abstract_id is not None and ilvl in abstract_lvl_defs.get(abstract_id, {}):
            return abstract_lvl_defs[abstract_id][ilvl]
        return ("decimal", "%1.")  # 查不到定義時的保守預設值

    return get_level_def

def normalize_and_solidify_list_numbering(doc: docx.Document):
    """
    將 Word 動態編號實體化成文字。

    重要修正（2026-08，第一輪）：
    編號來源分兩種——直接掛在段落自己的 w:numPr，或是繼承自段落套用的
    「樣式」(styles.xml 裡該 style 的 w:numPr)。過去版本在「繼承自樣式」
    的情況下，抓到的 numPr 節點其實是樣式本身共用的節點，處理完後又把它
    整個刪除，等於直接破壞了樣式定義——導致同一個樣式底下，只有第一個
    段落被正確編號，後面所有共用該樣式的段落全部失去編號（因為樣式的
    numPr 已經被上一個段落刪光了）。
    現在改用 get_effective_numpr 統一查找（不會動到樣式本身），並且
    「取消編號」一律透過在段落自己身上明確寫入 numId="0" 來完成
    （OOXML 規範：段落自己的 numId=0 代表關閉編號，且只影響這一個段落，
    不會動到樣式或其他共用該樣式的段落）。

    重要修正（2026-08，第二輪）：
    每個 numId 各自代表一份「獨立」的清單，彼此的計數應該互不影響。
    舊版有一段「切換到不同 numId 且在第 0 層時，就把 level_counters 整個
    清空」的邏輯，原意可能是想在切換清單時重置計數，但它是「不分青紅皂白
    清空全部 numId 的計數」，一旦文件中間夾雜了一份不相關的清單（例如本例
    中，標題自動編號 numId=14 的中間，插入了一份項目符號清單 numId=8），
    就會把完全不相關的 numId=14 的計數也一併洗掉，導致後續編號從頭重算、
    整個跑掉（例如子項目重複編號、父層級的編號不會往下遞增）。
    現在拿掉這段「跨 numId 全部清空」的邏輯：每個 (numId, ilvl) 的計數器
    完全獨立累加，只有在「同一個 numId」內，出現較淺層級的段落時，才清掉
    該 numId 自己「更深層」的計數（這部分維持原本邏輯，是正確的）。
    """
    get_level_def = _build_numbering_level_defs(doc)

    # 記錄各 numId 當前的階層計數：(numId, ilvl) -> count
    # 每個 numId 完全獨立，不會因為文件中出現了別的 numId 而被互相影響。
    level_counters = {}

    processed_count = 0

    for p in doc.paragraphs:
        result = get_effective_numpr(doc, p)

        if result is not None:
            num_id, ilvl, is_direct = result
            ilvl_int = int(ilvl)

            # 清理「同一個 numId」底下、比目前層級更深的子層計數
            # （例如從 1.1.3 回到 1.2 時，要把原本 ilvl=2 的計數清掉，
            #  下次再進到 ilvl=2 才會從 1 重新算起；但完全不影響其他 numId）
            keys_to_del = [k for k in level_counters.keys() if k[0] == num_id and k[1] > ilvl_int]
            for k in keys_to_del:
                del level_counters[k]

            own_fmt, own_lvl_text = get_level_def(num_id, ilvl)
            counter_key = (num_id, ilvl_int)

            if own_fmt == "bullet":
                prefix = "• "
            else:
                level_counters[counter_key] = level_counters.get(counter_key, 0) + 1

                def _replace_placeholder(m, current_num_id=num_id):
                    ref_ilvl_int = int(m.group(1)) - 1
                    ref_count = level_counters.get((current_num_id, ref_ilvl_int), 1)
                    ref_fmt, _ = get_level_def(current_num_id, str(ref_ilvl_int))
                    return _format_counter(ref_count, ref_fmt)

                prefix = re.sub(r'%(\d)', _replace_placeholder, own_lvl_text) + " "

            text_content = p.text.strip()

            # 放寬正則判斷，確保動態編號能順利寫入段落最前面
            if text_content:
                if p.runs:
                    # 清掉第一個 run 開頭原本就存在的空白（例如使用者手滑多打的
                    # 前導空格），避免跟編號後面自帶的空格疊成雙空格。
                    # 只去掉最前面的空白字元本身，不影響後面文字內容。
                    p.runs[0].text = prefix + p.runs[0].text.lstrip(' \u3000')
                else:
                    new_run = p.add_run(prefix)
                    p._p.remove(new_run._r)
                    p._p.insert(0, new_run._r)

                processed_count += 1

            # 取消這個段落的編號：
            # - 若編號是直接掛在段落自己的 w:numPr，直接移除該節點即可。
            # - 若編號是繼承自樣式，段落自己並沒有 w:numPr 節點可刪，
            #   必須明確新增一個 numId="0" 的 w:numPr 覆蓋掉繼承的編號，
            #   絕不能去動樣式本身的 w:numPr（那是共用資源）。
            pPr = p._p.get_or_add_pPr()
            existing_numPr = pPr.find(qn('w:numPr'))
            if is_direct:
                if existing_numPr is not None:
                    pPr.remove(existing_numPr)
            else:
                if existing_numPr is not None:
                    pPr.remove(existing_numPr)
                override_numPr = pPr.makeelement(qn('w:numPr'), {})
                override_numId = pPr.makeelement(qn('w:numId'), {qn('w:val'): '0'})
                override_numPr.append(override_numId)
                pPr.insert(0, override_numPr)

    print(f"[編號硬寫] 成功將 {processed_count} 個自動編號段落實體化為純文字")

def convert_docx_to_pdf(docx_path: Path, output_pdf_path: Path) -> Path:
    print(f"[PDF轉檔] 開始轉換 PDF: {docx_path.name}")
    try:
        if platform.system() == "Windows":
            if not os.path.exists(LIBREOFFICE_WIN_PATH):
                raise FileNotFoundError(
                    f"找不到 LibreOffice 執行檔，請確認已安裝於: {LIBREOFFICE_WIN_PATH}"
                )
            cmd = [
                LIBREOFFICE_WIN_PATH,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_pdf_path.parent.resolve()),
                str(docx_path.resolve())
            ]
        else:
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_pdf_path.parent.resolve()),
                str(docx_path.resolve())
            ]

        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        generated_pdf = output_pdf_path.parent / f"{docx_path.stem}.pdf"
        if generated_pdf.exists() and generated_pdf != output_pdf_path:
            generated_pdf.replace(output_pdf_path)

        print(f"[PDF轉檔] 轉換成功: {output_pdf_path.name}")
        return output_pdf_path
    except Exception as e:
        print(f"[PDF轉檔] 發生錯誤: {str(e)}")
        raise RuntimeError(f"LibreOffice 轉檔失敗: {str(e)}")


def delete_existing_document_and_clean_local(dataset_id: str, original_filename: str, headers: dict) -> bool:
    """
    查詢 Dify 知識庫中是否存在同名文件。
    若存在，會：
      1. 呼叫 Dify 刪除文件 API，把舊文件從知識庫移除
      2. 解析其舊 UUID 並刪除本地 docx / pdf 兩個資料夾對應的舊實體檔案
    回傳是否有刪除到舊文件（True/False）。

    改為「刪除舊文件＋一律新增」而非「update-by-file」，
    原因：update-by-file 無法變更既有文件的 doc_form（一般 / 父子模式），
    若舊文件是在改用 hierarchical 模式之前建立的，更新時會被 Dify 拒絕
    （錯誤：doc_form is different from the dataset doc_form）。
    刪除後重新建立，可確保每次都是乾淨的 hierarchical 模式文件。

    嚴謹處理：若查詢文件列表失敗、或找到同名文件卻刪除失敗，
    一律直接丟出例外中止整個上傳流程，避免 Dify 裡同時存在
    「刪不掉的舊文件」與「新建立的文件」兩份同名重複文件。
    """
    print(f"[狀態查詢] 檢查 Dify 是否已有同名文件: '{original_filename}'")
    list_url = f"{DIFY_API_BASE}/datasets/{dataset_id}/documents"
    params = {"page": 1, "limit": 100}
    response = requests.get(list_url, headers=headers, params=params)

    if response.status_code != 200:
        raise Exception(
            f"查詢知識庫文件列表失敗，狀態碼: {response.status_code}，內容: {response.text}"
        )

    data = response.json()
    documents = data.get("data", [])

    for doc in documents:
        if doc.get("name") == original_filename:
            doc_id = doc.get("id")
            print(f"[狀態查詢] 找到舊版本文件，ID: {doc_id}，開始清理...")

            # 先刪除 Dify 上的舊文件；失敗就直接中止，不繼續往下建立新文件，
            # 避免同一份檔名在知識庫裡同時存在新舊兩份。
            delete_url = f"{DIFY_API_BASE}/datasets/{dataset_id}/documents/{doc_id}"
            del_res = requests.delete(delete_url, headers=headers)
            if del_res.status_code in (200, 204):
                print(f"[Dify刪除] 已成功刪除舊文件，ID: {doc_id}")
            else:
                raise Exception(
                    f"刪除舊文件失敗（ID: {doc_id}），狀態碼: {del_res.status_code}，"
                    f"內容: {del_res.text}，已中止上傳流程以避免產生重複文件。"
                )

            # 優先嘗試：直接從文件列表回傳的 metadata 解析舊 UUID
            old_uuid = get_file_uuid_from_metadata(doc)

            if old_uuid:
                print(f"[狀態查詢] 從 metadata 解析出舊檔案 UUID: {old_uuid}")
                safe_unlink(PDF_DIR / f"{old_uuid}.pdf")
                safe_unlink(DOCX_DIR / f"{old_uuid}.docx")
            else:
                # 退回舊機制：從 segment 內文用 regex 解析 PDF 網址取得 UUID
                # （適用於舊版尚未寫入 metadata 的歷史文件）
                print("[狀態查詢] metadata 中未找到 file_uuid，退回舊的內文解析機制。")
                segments_url = f"{DIFY_API_BASE}/datasets/{dataset_id}/documents/{doc_id}/segments"
                seg_res = requests.get(segments_url, headers=headers)

                if seg_res.status_code == 200:
                    seg_data = seg_res.json().get("data", [])
                    if seg_data:
                        first_chunk_text = seg_data[0].get("content", "")
                        match = re.search(r'/static/pdf/([a-f0-9\-]+\.pdf)', first_chunk_text)
                        if match:
                            old_pdf_filename = match.group(1)
                            old_uuid = Path(old_pdf_filename).stem
                            print(f"[狀態查詢] 解析出舊檔案 UUID: {old_uuid}")
                            safe_unlink(PDF_DIR / old_pdf_filename)
                            safe_unlink(DOCX_DIR / f"{old_uuid}.docx")
                        else:
                            print("[狀態查詢] 警告：無法從舊文件的 segment 內容解析出 UUID。")

            return True

    print("[狀態查詢] 知識庫中未找到同名文件，將執行新增流程。")
    return False


async def wait_for_dify_indexing(dataset_id: str, document_id: str, headers: dict, max_retries: int = 60, delay: int = 5) -> dict:
    print(f"[狀態輪詢] 開始輪詢文件索引狀態，文件 ID: {document_id}")
    for i in range(max_retries):
        url = f"{DIFY_API_BASE}/datasets/{dataset_id}/documents/{document_id}"
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            doc_info = response.json()
            doc_data = doc_info.get("document", doc_info)
            status = doc_data.get("indexing_status")

            print(f"[狀態輪詢] 第 {i+1} 次檢查，目前狀態: {status}")

            if status == "completed":
                print("[狀態輪詢] 索引完成！")
                return doc_info
            elif status == "error":
                print("[狀態輪詢] 索引失敗，Dify 返回錯誤狀態。")
                return None
        else:
            print(f"[狀態輪詢] 請求失敗，狀態碼: {response.status_code}")

        await asyncio.sleep(delay)

    print("[狀態輪詢] 達到最大重試次數，索引發生超時。")
    return None


@app.post("/upload-to-dify")
async def process_and_upload_to_dify(file: UploadFile = File(...), dataset_id: str = Form(...)):
    if not file.filename.endswith(".docx"):
        return {
            "status": "error",
            "detail": "僅支援 .docx 格式檔案"
        }

    original_filename = file.filename
    print(f"\n========== 開始處理上傳請求: {original_filename}（知識庫: {dataset_id}）==========")
    headers = {"Authorization": f"Bearer {DIFY_API_KEY}"}

    file_uuid = str(uuid.uuid4())
    print(f"[初始化] 生成新檔案 UUID: {file_uuid}")

    docx_path = DOCX_DIR / f"{file_uuid}.docx"
    pdf_path = PDF_DIR / f"{file_uuid}.pdf"

    try:
        file_bytes = await file.read()
        with open(docx_path, "wb") as buffer:
            buffer.write(file_bytes)
        print("[儲存檔案] 使用者上傳檔已寫入本地儲存區，準備進行處理。")

        # 依序套用兩個處理，處理完直接存回同一份 docx_path：
        # 1. 縮排修正：把編號段落該有的縮排值寫死到段落自己身上（供 LibreOffice 正確排版）
        # 2. 編號硬寫：把 1. / 1.1 / 1.1.1 這種自動編號實體化成文字，並移除 w:numPr，
        #    確保 Dify 抽取文字時看得到編號（順序必須在縮排修正之後，理由見函式註解）
        # 這份處理後的文件會同時拿去轉 PDF 跟餵給 Dify，兩邊看到的編號內容完全一致。
        doc_for_processing = docx.Document(docx_path)
        fix_list_indentation_for_libreoffice(doc_for_processing)
        normalize_and_solidify_list_numbering(doc_for_processing)
        doc_for_processing.save(docx_path)
        del doc_for_processing
        gc.collect()
        print("[文件處理] 縮排修正與編號硬寫皆已完成並存回。")

        convert_docx_to_pdf(docx_path, pdf_path)

        pdf_url = f"{SERVER_BASE_URL}/static/pdf/{file_uuid}.pdf"

        # 若知識庫中已有同名文件，先刪除舊文件並清掉對應的本地舊檔案，
        # 之後一律走「建立新文件」流程（不再用 update-by-file，
        # 因為 update-by-file 無法變更既有文件的 doc_form）。
        delete_existing_document_and_clean_local(dataset_id, original_filename, headers)

        data = {
            "process_rule": {
                "mode": "hierarchical",
                "rules": {
                    "pre_processing_rules": [
                        {"id": "remove_extra_spaces", "enabled": True},
                        {"id": "remove_urls_emails", "enabled": False}
                    ],
                    # 沿用文件裡使用者手動打的 ===CHUNK=== 標記來切分 parent chunk
                    # （這不是程式插入的，是原始文件內容本身就含有的分段標記）。
                    "segmentation": {
                        "separator": "===CHUNK===",
                        "max_tokens": 4000
                    },
                    "parent_mode": "paragraph",
                    "subchunk_segmentation": {
                        "separator": "\n\n",
                        "max_tokens": 4000
                    }
                }
            },
            "indexing_technique": "high_quality",
            "doc_form": "hierarchical_model",
            "retrieval_model": {
                "search_method": "semantic_search",
                "top_k": 3,
                "reranking_enable": False,
                "score_threshold_enabled": False
            }
        }

        with open(docx_path, "rb") as docx_file:
            files = {
                "file": (original_filename, docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                "data": (None, json.dumps(data), "application/json")
            }

            print("[Dify上傳] 開始呼叫 API 建立新文件（使用縮排修正＋編號硬寫後的版本）")
            api_url = f"{DIFY_API_BASE}/datasets/{dataset_id}/document/create-by-file"

            response = requests.post(api_url, headers=headers, files=files)

        if response.status_code != 200:
            print(f"[API失敗] 回應內容: {response.text}")
            raise Exception(f"Dify API 操作失敗: {response.text}")

        dify_response_data = response.json()
        target_doc_id = dify_response_data.get("document", {}).get("id")

        if not target_doc_id:
            raise Exception("無法從 Dify 回應中取得文件 ID")

        # 文件建立成功後，把這次的 file_uuid、pdf_url 一起寫入該文件的 metadata，
        # 之後要對照本地 docx / pdf 資料夾的檔案，就不用再靠內文 regex 解析。
        set_document_metadata(dataset_id, target_doc_id, file_uuid, pdf_url, headers)

        final_dify_response = await wait_for_dify_indexing(dataset_id, target_doc_id, headers)

        if final_dify_response:
            print("========== 處理流程全數成功完成 ==========\n")
            return {
                "status": "success",
                "original_filename": original_filename,
                "dataset_id": dataset_id,
                "pdf_url": pdf_url,
                "dify_response": final_dify_response
            }
        else:
            # 索引失敗：主動呼叫 Dify API 刪除這筆失敗的文件，避免殘留
            print(f"[清理機制] 索引失敗，準備從 Dify 知識庫移除文件 ID: {target_doc_id}")
            delete_url = f"{DIFY_API_BASE}/datasets/{dataset_id}/documents/{target_doc_id}"
            try:
                del_res = requests.delete(delete_url, headers=headers)
                if del_res.status_code in (200, 204):
                    print(f"[清理機制] 已成功從 Dify 知識庫移除失敗文件: {target_doc_id}")
                else:
                    print(f"[清理機制] 從 Dify 移除失敗文件失敗，狀態碼: {del_res.status_code}")
            except Exception as clean_err:
                print(f"[清理機制] 移除文件時發生例外: {str(clean_err)}")

            raise Exception("Dify 索引失敗或發生超時錯誤")

    except Exception as e:
        print(f"\n[錯誤處理] 流程發生例外，清理本次產生的暫存檔，錯誤內容: {str(e)}")
        safe_unlink(docx_path)
        safe_unlink(pdf_path)
        print("========== 錯誤處理執行完畢 ==========\n")
        return {
            "status": "error",
            "detail": str(e)
        }